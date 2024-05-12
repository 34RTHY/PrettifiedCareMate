import streamlit as st
from streamlit_option_menu import option_menu
from streamlit_extras.switch_page_button import switch_page 
import time
import pymongo
import os
from llama_index.vector_stores.mongodb import MongoDBAtlasVectorSearch
from llama_index.core import VectorStoreIndex
from llama_index.llms.openai import OpenAI
from llama_index.core.node_parser import SentenceSplitter
from llama_index.llms.together import TogetherLLM
from llama_index.core import Document
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import Settings


import os
from datetime import datetime
import PyPDF2
import docx
import io
import streamlit as st
import src.CareMate as caremate

path_to_css = './style/Appstyle.css'

class Chatapp:
    def __init__(self):
        st.set_page_config(initial_sidebar_state="auto")
        caremate.initialize.css_loader(path_to_css)

    def extract_text_from_txt(self,file_contents):
        return file_contents.decode("utf-8")

    def extract_text_from_docx(self,file_contents):
        doc = docx.Document(io.BytesIO(file_contents))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def extract_text_from_pdf(self,file_contents):
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_contents))
        text = ""
        for page_number in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_number].extract_text()
        return text
    
    def cached_extract_text(self,file_contents, file_type):
        if file_type == "text/plain":
            return self.extract_text_from_txt(file_contents)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self.extract_text_from_docx(file_contents)
        elif file_type == "application/pdf":
            return self.extract_text_from_pdf(file_contents)
        else:
            st.error("Unsupported file format.")

    @st.experimental_fragment
    def show_download_button(self,Label,data,file_name):
        st.download_button(label=Label,
                            data=data,
                            file_name=file_name,
                            )

    def maketopatient(self,text):
        parts = text.split('\n\n')
        str1 = parts[0].replace('## Possible diseases based on the symptoms described:','')
        str2 = parts[1].replace('## Treatments for each disease:','')
        str3 = parts[2].replace('## Specify whether the patient should go to a doctor or pharmacy:','')
        str4 = parts[3].replace('## Next steps for the patient:','')
        return str1.replace('*','').strip(), str2.replace('*','').strip(), str3.replace('*','').strip(), str4.replace('*','').strip()
    
    def maketodoctor(self,text):
        parts = text.split('\n\n')
        str1 = parts[0].replace('## Suggested Diagnosis:','')
        str2 = parts[1].replace('## Suggested Actions to Assist the Doctor:','')
        str3 = parts[2].replace('## Suggested Laboratory Tests for Precise Diagnosis:','')
        return str1.replace('*','').strip(), str2.replace('*','').strip(), str3.replace('*','').strip()

    def create_query_str(self,text) -> str:
        query_str = f"""
As a medical coding and billing tool,
you'll analyze clinical reports to suggest ICD, CPT, and HCPCS codes
for various medical terms, symptoms, diagnoses, treatments, and procedures.
Provide multiple relevant codes when necessary.

Here is the task: {text}

-----------Response Format-----------

\t- ICD 11 code: original text | code (code name) *Add more codes if relevant
\t- CPT code: original text | code (code name) *Add more codes if relevant
\t- HCPCS code: original text | code (code name) *Add more codes if relevant
*Add more Medical Terminology if relevent

-----------Response Format-----------

"""
        return query_str
    
    def app(self):
        invalid_chars = ''' :/'*\\?"<>|'''
        with st.sidebar:
            selected=option_menu(
                menu_title='CareMate Mode Menu',
                options=['Medical Coding','Patient Mode','Doctor Mode'],
                menu_icon='heart-eyes-fill',
                default_index=0,
                styles={
                    "container": {"padding": "5!important","background-color":'#fafaf4',"box-shadow": "0 2px 4px rgba(0, 0, 0, 0.1)"},
                    "icon": {"color": "#0f162a", "font-size": "23px"}, 
                    "nav-link": {"color":"#0f162a","font-size": "20px", "text-align": "left", "margin":"0px", "--hover-color": "#b7caff","border-radius": "17px","outline": "2px solid #0f162a"},
                    "nav-link-selected": {"background-color": "#0f162a","color":"white"},}
            )

            if st.button("Back",type='primary'):
                switch_page('Home')

        Settings.embed_model = OpenAIEmbedding(model_name='text-embedding-3-small')
        gpt3_5 = OpenAI(model="gpt-3.5-turbo-0125", temperature=0,api_key =  st.secrets["OPENAI_API_KEY"])
        dbrx = TogetherLLM(model="databricks/dbrx-instruct", api_key=st.secrets['TOGETHER_API_KEY'])

        client = pymongo.MongoClient(st.secrets["MongoUri"])
        clinet_cb = pymongo.MongoClient(st.secrets["client_cb"])
        ### Database
        #Store_cb
        store_cb = MongoDBAtlasVectorSearch(
                                    mongodb_client = clinet_cb,
                                    db_name = 'CareMate',
                                    collection_name = 'Medical-Coding-And-Billing',
                                    index_name = 'vector_index',
                                )
        index_cb = VectorStoreIndex.from_vector_store(store_cb)
        #Store
        store = MongoDBAtlasVectorSearch(
                                    mongodb_client = client,
                                    db_name = 'Embeddings',
                                    collection_name = '3small',
                                    index_name = 'vector_index',
                                    embedding_key = '3small',
                                )
        index = VectorStoreIndex.from_vector_store(store)

        ##Base query engine_cb
        base_query_engine_cb = index_cb.as_query_engine(
            similarity_top_k=5,
            node_postprocessors=[caremate.Carellm.load_reranker(3,gpt3_5)],
            llm = dbrx,
            temperature = 0
        )

        ##Base query engine
        base_query_engine = index.as_query_engine(
            similarity_top_k=12,
            node_postprocessors=[caremate.Carellm.load_reranker(5,gpt3_5)],
            llm = dbrx,
            temperature = 0
        )
        modified_query_engine = caremate.Carellm.load_query_transform_engine(base_query_engine,dbrx)


        ###App

        if selected=='Medical Coding':
            st.title('Medical Coding Using :blue[CareMate]')
            write_result = False
            with st.container(border = True):
                uploaded_file = st.file_uploader("Choose a file",type = ['docx','pdf','txt'])
                if uploaded_file is not None:
                    if st.button('Medical Coding',type='primary'):
                        generated_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        write_result = True
            if write_result:
                #upload file
                file_contents = uploaded_file.getvalue()
                file_type = uploaded_file.type

                text = self.cached_extract_text(file_contents, file_type)
                doc = [Document(text=text)]
                splitter = SentenceSplitter(
                    chunk_size=128,
                    chunk_overlap=10,
                )
                nodes = splitter.get_nodes_from_documents(doc, show_progress = False)

                ###Progressbar
                progress_text = "Operation in progress. Please wait."
                progression = 0
                my_bar = st.empty()
                my_bar = st.progress(progression, text=progress_text)

                ###Querry
                suggested_medical_codes = []
                for node in nodes:
                    suggested_medical_codes.append(base_query_engine_cb.query(self.create_query_str(str(node.text))))
                    progression += 100//len(nodes)
                    my_bar.progress(progression, text=progress_text)

                suggested_medical_codes = [i.response for i in suggested_medical_codes]
                unsorted_suggested_medical_codes = '\n'.join(suggested_medical_codes)

                ICD_code_list = []
                CPT_code_list = []
                HCPCS_code_list = []
                for i in unsorted_suggested_medical_codes.split('\n'):
                    tmp_string = i.replace('*Add more codes if relevant', '')
                    try:
                        if tmp_string[0] == '-' and 'No specific' not in tmp_string and 'Not specified' not in tmp_string:
                            if tmp_string[2] == 'I':
                                ICD_code_list.append(tmp_string)
                            elif tmp_string[2] == 'C':
                                CPT_code_list.append(tmp_string)
                            elif tmp_string[2] == 'H':
                                HCPCS_code_list.append(tmp_string)
                    except:
                        continue

                Suggested_ICD_11_Codes_Text = "\n".join([ '\t' + "".join(i.split(': ')[1:]) for i in ICD_code_list])
                Suggested_CPT_Codes_Text = "\n".join([ '\t' + "".join(i.split(': ')[1:]) for i in CPT_code_list])
                Suggested_HCPCS_Codes_Text = "\n".join([ '\t' + "".join(i.split(': ')[1:]) for i in HCPCS_code_list])

                my_bar.progress(100, text='Completed')
                time.sleep(1)
                ###TOPDF
                filename = ''.join(c if c not in invalid_chars else '_' for c in uploaded_file.name)
                timestamp = ''.join(c if c not in invalid_chars else '_' for c in generated_time)
                File_stamp = f"{filename}_{timestamp}"
                caremate.writer.generate_medical_code_report(Suggested_ICD_11_Codes_Text,Suggested_CPT_Codes_Text,Suggested_HCPCS_Codes_Text,f'./pdfout/{File_stamp}.pdf')

                with open(f"./pdfout/{File_stamp}.pdf", "rb") as pdf_file:
                    PDFbyte = pdf_file.read()

                os.remove(f"./pdfout/{File_stamp}.pdf")

                self.show_download_button("Download Report",PDFbyte,"Medical Code Report.pdf")


                ###OUTPUT
                with st.container(border = True):
                    st.markdown('## Suggested ICD 11 Codes')
                    for Each_code in Suggested_ICD_11_Codes_Text.split('\t'):
                        if Each_code.strip() == '' or Each_code == None or None:
                            continue
                        str_list = Each_code.split('|')
                        Symptom_name = str_list[0].strip()
                        try:
                            Code_name = str_list[1].strip()
                        except:
                            Code_name = str_list[0].strip()
                        st.markdown(f'- **{Symptom_name}** : {Code_name}')
                with st.container(border = True):
                    st.markdown('## Suggested CPT Codes')
                    for Each_code in Suggested_CPT_Codes_Text.split('\t'):
                        if Each_code.strip() == '' or Each_code == None or None:
                            continue
                        str_list = Each_code.split('|')
                        Symptom_name = str_list[0].strip()
                        try:
                            Code_name = str_list[1].strip()
                        except:
                            Code_name = str_list[0].strip()
                        st.markdown(f'- **{Symptom_name}** : {Code_name}')
                with st.container(border = True):
                    st.markdown('## Suggested HCPCS Codes')
                    for Each_code in Suggested_HCPCS_Codes_Text.split('\t'):
                        if Each_code.strip() == '' or Each_code == None or None:
                            continue
                        str_list = Each_code.split('|')
                        Symptom_name = str_list[0].strip()
                        try:
                            Code_name = str_list[1].strip()
                        except:
                            Code_name = str_list[0].strip()
                        st.markdown(f'- **{Symptom_name}** : {Code_name}')

        if selected=='Patient Mode':
            Askpatient = False
            st.title('Ask :blue[CareMate] to analyze your symptoms or injuries')
            with st.container(border = True):
                AgeNumber = st.number_input('Insert your Age',min_value=0)

                Medical_History = st.text_area(
                    "Medical History",
                    help='''Any pre-existing medical conditions (e.g., diabetes, hypertension), Current medications (prescription and over-the-counter), Allergies (including medication, food, and environmental allergies), Previous surgeries or hospitalizations''',
                    key = 'MedHis',
                    placeholder='''I'm a cancer survivor with a tumour on my left arm. I received treatment, and now seeking follow-up care. There are no visible symptoms to be shown'''
                )

                Symptoms_injuries = st.text_area(
                    "Symptoms/Injuries details",
                    help='''Description of symptoms (including onset, duration, severity, and any triggers), Any accompanying symptoms (e.g., fever, nausea, dizziness), Changes in symptoms over time, Nature of the injury (e.g., blunt trauma, laceration, burn), Mechanism of injury (how the injury occurred), Any loss of consciousness or impact on consciousness''',
                    placeholder='I am experiencing severe cramps in my left arm whenever I attempt to bend over to pick something up from the ground. The pain is intense and makes it challenging to perform this action.',
                )

                timeline = st.text_area(
                    "Timeline",
                    help='''When did the symptoms or injuries first appear, and have they changed or progressed over time?''',
                    placeholder = 'I began experiencing sharp pain in my left shoulder three days ago. The pain is persistent and localized to the left shoulder area.',
                )

                Additional_Context_and_Family_History = st.text_area(
                    "Additional Context and Family History",
                    help='''Share any relevant lifestyle factors, recent travel history, occupation, family medical history, or exposure to sick individuals.''',
                    placeholder='I maintain a balanced diet, exercise regularly, and avoid smoking or recreational drugs. I occasionally have a glass of wine with dinner.I also recently traveled to Europe for two weeks without experiencing any illness. I followed COVID-19 safety guidelines. Family history includes heart disease in parents and grandparents, diabetes in one grandparent, and my sister had breast cancer last year.'
                )
                if st.button('Ask CareMate',type = 'primary'):
                    Askpatient = True
            

            if Askpatient:
                generated_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                Patient_input = f'My age:{AgeNumber}, My medical history:{Medical_History}, My symptoms:{Symptoms_injuries}, The timeline of this symptoms:{timeline} and Additional Context and Family History:{Additional_Context_and_Family_History}'
                query_str = f"""As a professional medical service, diagnose three possible diseases based on the provided data:{Patient_input}. Use bullet points to list diseases from most likely to least likely. Include reasons for each diagnosis.

If unsure, avoid sharing false information. Outline treatment options for each identified disease and specify whether the patient should seek professional medical attention or opt for self-care at a pharmacy.

Response in MarkDown Format:
## Possible diseases based on the symptoms described:
- **Name Of Disease 1** : Reason
- **Name Of Disease 2** : Reason
- **Name Of Disease 3** : Reason

## Treatments for each disease:
- **Name Of Disease 1** : Treatment 1
- **Name Of Disease 2** : Treatment 2
- **Name Of Disease 3** : Treatment 3


## Specify whether the patient should go to a doctor or pharmacy:
- **Answer**: Reason

## Next steps for the patient:
- **Answer** : Reason
"""


                        
                with st.spinner(text="In progress..."):
                    response = modified_query_engine.query(query_str)
                full_response = response.response

                Possible_diseases, Treatments_for_each_disease,doctor_or_pharmacy,Next_steps_for_the_patient = self.maketopatient(full_response)

                timestamp = ''.join(c if c not in invalid_chars else '_' for c in generated_time)
                File_stamp = f"{timestamp}"
                caremate.writer.generate_patient_mode_report(Possible_diseases, Treatments_for_each_disease,doctor_or_pharmacy,Next_steps_for_the_patient,f'./pdfout/{File_stamp}.pdf')
                with open(f"./pdfout/{File_stamp}.pdf", "rb") as pdf_file:
                    PDFbyte = pdf_file.read()

                os.remove(f"./pdfout/{File_stamp}.pdf")

                self.show_download_button("Download Report",PDFbyte,"Patient Mode Report.pdf")
                        
                with st.container(border = True):
                    st.markdown(full_response)

                    


        if selected=='Doctor Mode':
            Askdoctor = False
            st.title('Ask :blue[CareMate] to analyze your symptoms or injuries')
            with st.container(border = True):
                task_input = st.text_area(
                    "Task Input",
                    help = '''Tasks or questions regarding diagnosing the disease''',
                    placeholder='''Please determine whether or not this patient has an HIV infection.'''
                )

                patient_description = st.text_area(
                    "Patient Description",
                    help = '''Patient description used for diagnosing the symptoms.''',
                    placeholder = '''Gender: Male
Age: 35
Occupation: Office administrator
Presenting Symptoms:
Persistent fever for the past two weeks
Generalized weakness and fatigue
Night sweats
Unintentional weight loss of 4 kg over the last month
Oral thrush
Chronic diarrhea for the past two weeks
History of recurrent herpes zoster infections
Social History:
Single, lives alone
Non-smoker, occasional alcohol consumption
Limited social support network
Compliance with ART regimen has been inconsistent due to work-related stress'''
                )
                Lab_result = st.text_area(
                    "Lab results",
                    help = '''Your Laboratory result''',
                    placeholder='''LYMPH SUBSET
CD3 ABS, RESULT:1920, UNITS:cu.mm, REFERENCE RANGE:625-2460
CD3 %, RESULT:82.0, UNITS:Percent, REFERENCE RANGE:60-90
CD3+/CD4+ (HELPER) ABS, RESULT:570, UNITS:cu.mm, REFERENCE RANGE:423-1724
CD3+/CD4+ (HELPER) %, RESULT: 23.8 L, UNITS:Percent, REFERENCE RANGE:32-68
CD3+/CD8+ (SUPPRES) ABS, RESULT:1290 H, UNITS:cu.mm, REFERENCE RANGE:140-958
CD3+/CD8+ (SUPPRES) %, RESULT:53.9 H UNITS:Percent, REFERENCE RANGE:10-36
CD4/CD8 RATIO, RESULT:0.44 L, UNITS:Ratio, REFERENCE RANGE:0.90-6.00'''
                )
                if st.button('Ask CareMate',type = 'primary',key = 'Doctor'):
                    Askdoctor = True
            
            if Askdoctor:
                generated_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                query_str = f"""As a professional medical health service provider, our task is to accurately assess possible diseases based on the given clinical laboratory test data and patient description, while ensuring the response remains faithful to the provided context. If the patient is healthy and unaffected, we will indicate that the patient's health is normal. We prioritize providing accurate information and refrain from sharing false information.

Patient Presentation:
The patient presents with the following clinical laboratory test data:
- Patient Description: {patient_description}
- Laboratory Test Result: {Lab_result}
- Question: {task_input}

Your task:
Please find below the possible diseases and reasons for their consideration, listed from most likely to least likely:
To assist the doctor in accurately diagnosing the symptoms, we recommend the following actions in priority order:
To enhance diagnostic accuracy, we suggest conducting the following laboratory tests in priority order:

-----RESPONSE in Markdown FORMAT-----

## Suggested Diagnosis:
- **Disease** : Reason
- **Disease** : Reason (if suspected)
- **Disease** : Reason (if suspected)

## Suggested Actions to Assist the Doctor:
- **Recommendation** : Reason
- **Recommendation** : Reason (if relevant)
- **Recommendation** : Reason (if relevant)

## Suggested Laboratory Tests for Precise Diagnosis:
- **Test** : Reason
- **Test** : Reason (if relevant)
- **Test** : Reason (if relevant)

-----RESPONSE FORMAT-----
"""

                                
                    
                with st.spinner(text="In progress..."):
                    response = modified_query_engine.query(query_str)
                full_response = response.response

                Suggested_Diagnosis, Suggested_Actions,Suggested_Laboratory_Tests = self.maketodoctor(full_response)

                timestamp = ''.join(c if c not in invalid_chars else '_' for c in generated_time)
                File_stamp = f"{timestamp}"
                caremate.writer.generate_doctor_mode_report(Suggested_Diagnosis, Suggested_Actions,Suggested_Laboratory_Tests,f'./pdfout/{File_stamp}.pdf')
                with open(f"./pdfout/{File_stamp}.pdf", "rb") as pdf_file:
                    PDFbyte = pdf_file.read()

                os.remove(f"./pdfout/{File_stamp}.pdf")

                self.show_download_button("Download Report",PDFbyte,"Doctor Mode Report.pdf")

                with st.container(border = True):
                    st.markdown(full_response)

Chat_instance = Chatapp()
Chat_instance.app()



